"""
document_parser.py
-------------------
Extracts raw text and paragraph-level segments from PDF, DOCX, and TXT files
so that downstream rule matching can operate on clean, normalized input.
"""

from __future__ import annotations

import os
import re
from dataclasses import dataclass, field
from typing import List


@dataclass
class ParsedDocument:
    filename: str
    full_text: str
    paragraphs: List[str] = field(default_factory=list)

    @property
    def word_count(self) -> int:
        return len(self.full_text.split())


class UnsupportedFileTypeError(Exception):
    pass


def _clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _split_paragraphs(text: str) -> List[str]:
    raw_paras = re.split(r"\n\s*\n", text)
    return [p.strip() for p in raw_paras if p.strip()]


def parse_txt(path: str) -> ParsedDocument:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        raw = f.read()
    cleaned = _clean_text(raw)
    return ParsedDocument(
        filename=os.path.basename(path),
        full_text=cleaned,
        paragraphs=_split_paragraphs(cleaned),
    )


def parse_docx(path: str) -> ParsedDocument:
    try:
        import docx  # python-docx
    except ImportError as e:
        raise ImportError(
            "python-docx is required to parse .docx files. "
            "Install with `pip install python-docx`."
        ) from e

    document = docx.Document(path)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    # Also pull text out of tables, since contracts often store terms in them
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    full_text = _clean_text("\n\n".join(paragraphs))
    return ParsedDocument(
        filename=os.path.basename(path),
        full_text=full_text,
        paragraphs=paragraphs,
    )


def parse_pdf(path: str) -> ParsedDocument:
    try:
        import pdfplumber
    except ImportError as e:
        raise ImportError(
            "pdfplumber is required to parse .pdf files. "
            "Install with `pip install pdfplumber`."
        ) from e

    all_text = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            all_text.append(page_text)

    full_text = _clean_text("\n\n".join(all_text))
    return ParsedDocument(
        filename=os.path.basename(path),
        full_text=full_text,
        paragraphs=_split_paragraphs(full_text),
    )


def parse_document(path: str) -> ParsedDocument:
    """Dispatch to the correct parser based on file extension."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".txt":
        return parse_txt(path)
    if ext == ".docx":
        return parse_docx(path)
    if ext == ".pdf":
        return parse_pdf(path)
    raise UnsupportedFileTypeError(f"Unsupported file type: {ext}")
